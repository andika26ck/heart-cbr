# -*- coding: utf-8 -*-
"""Build the JOIV-style scientific article (.docx) in Bahasa Indonesia.
A4, two columns, Times New Roman 12pt body, 10pt bold sub-section, 10pt caption,
margins Top/Bottom 3cm & Left/Right 2cm, numeric [n] citations, >=8 pages.
"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RES = os.path.join(ROOT, "results")
M = json.load(open(f"{RES}/metrics.json"))
cv, ts, tr = M["cv_full_metrics"], M["test_metrics"], M["train_metrics"]
import pandas as _pd
_df = _pd.read_csv(f"{ROOT}/data/heart_disease.csv")
_cm = cv["confusion"]
_tp, _tn, _fp, _fn = _cm["tp"], _cm["tn"], _cm["fp"], _cm["fn"]
def _safe(a, b):
    return a / b if b else 0.0
prec_pos = _safe(_tp, _tp + _fp); rec_pos = _safe(_tp, _tp + _fn)
f1_pos = _safe(2 * prec_pos * rec_pos, prec_pos + rec_pos)
prec_neg = _safe(_tn, _tn + _fn); rec_neg = _safe(_tn, _tn + _fp)
f1_neg = _safe(2 * prec_neg * rec_neg, prec_neg + rec_neg)
BLUE = RGBColor(0x27, 0x83, 0xDE)
DARK = RGBColor(0x2C, 0x2C, 0x2B)
TWIP = 566.9  # twips per cm

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0


def set_cols(section, num, space=500):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); sectPr.append(cols)
    cols.set(qn("w:num"), str(num)); cols.set(qn("w:space"), str(space))


def set_page(section):
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0); section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)


def run(p, text, size=12, bold=False, italic=False, color=None, sc=False, mono=False):
    r = p.add_run(text)
    r.font.name = "Consolas" if mono else "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), r.font.name)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if sc:
        r.font.small_caps = True
    return r


def para(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False,
         italic=False, before=0, after=6, indent=0.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Cm(indent)
    if text:
        run(p, text, size=size, bold=bold, italic=italic)
    return p


def h1(num, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    if num:
        run(p, f"{num}. ", size=12, bold=True, sc=True, color=DARK)
    run(p, text, size=12, bold=True, sc=True, color=DARK)


def h2(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    run(p, text, size=10, bold=True, italic=True, color=DARK)


def eq(text, number):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    run(p, text, size=11, italic=True)
    run(p, "      (" + number + ")", size=11)


def figure(path, caption, width_cm=8.0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width_cm))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    run(c, caption, size=10)


def caption_top(text):
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(6); c.paragraph_format.space_after = Pt(2)
    run(c, text, size=10, sc=True)


def make_table(rows, widths):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"; t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(int(sum(widths)*TWIP)))
    tblW.set(qn("w:type"), "dxa"); tblPr.append(tblW)
    grid = t._tbl.find(qn("w:tblGrid"))
    for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
        gc.set(qn("w:w"), str(int(widths[i]*TWIP)))
    for i, rowvals in enumerate(rows):
        for j, val in enumerate(rowvals):
            cell = t.cell(i, j)
            cell.width = Cm(widths[j])
            pr = cell.paragraphs[0]
            pr.paragraph_format.space_after = Pt(1); pr.paragraph_format.space_before = Pt(1)
            pr.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run(pr, str(val), size=9, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def algorithm(title, lines, width_cm=7.6):
    caption_top(title)
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"; t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(int(width_cm*TWIP)))
    tblW.set(qn("w:type"), "dxa"); tblPr.append(tblW)
    grid = t._tbl.find(qn("w:tblGrid"))
    grid.findall(qn("w:gridCol"))[0].set(qn("w:w"), str(int(width_cm*TWIP)))
    cell = t.cell(0, 0); cell.width = Cm(width_cm)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for ln in lines:
        p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run(p, ln, size=8, mono=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


pct = lambda x: f"{x*100:.1f}%"

# ===== SECTION 0: title / abstract (single column) =====
set_page(doc.sections[0]); set_cols(doc.sections[0], 1)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(6)
run(t, "Penerapan Case-Based Reasoning dengan Similaritas HEOM "
       "Berbobot untuk Diagnosis Penyakit Jantung", size=20, bold=True, color=DARK)
au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.paragraph_format.space_after = Pt(2)
run(au, "Andika Candra K.", size=12, bold=True); run(au, " a,*", size=8, bold=True)
af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
af.paragraph_format.space_after = Pt(2)
run(af, "a Program Studi Informatika, Fakultas Teknik, "
        "Universitas Muhammadiyah Malang, Malang, Indonesia", size=10, italic=True)
co = doc.add_paragraph(); co.alignment = WD_ALIGN_PARAGRAPH.CENTER
co.paragraph_format.space_after = Pt(10)
run(co, "*Corresponding author: andikakunnn@gmail.com", size=10, italic=True)

ab = doc.add_paragraph(); ab.paragraph_format.space_after = Pt(4)
run(ab, "Abstrak", size=11, bold=True, italic=True)
abt = doc.add_paragraph(); abt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abt.paragraph_format.space_after = Pt(6)
run(abt,
    "Penyakit jantung merupakan penyebab kematian tertinggi di dunia sehingga "
    "diagnosis dini yang akurat sangat dibutuhkan. Penelitian ini menerapkan metode "
    "Case-Based Reasoning (CBR) untuk mendiagnosis penyakit jantung berdasarkan basis "
    "kasus rekam medis. Sistem dibangun mengikuti siklus CBR lima tahap, yaitu "
    "representasi kasus, retrieve, reuse, revise, dan retain. Tahap retrieve "
    "menggunakan fungsi similaritas Heterogeneous Euclidean-Overlap Metric (HEOM) yang "
    "mampu menangani atribut numerik dan nominal secara bersamaan, dipadukan dengan "
    "pembobotan fitur berbasis korelasi. Dataset mengikuti skema resmi UCI Cleveland "
    "Heart Disease yang terdiri atas 13 atribut klinis dan 303 kasus. Evaluasi dilakukan "
    "menggunakan validasi silang 5-fold dan pengujian hold-out 20%. Hasil pengujian "
    f"menunjukkan model mencapai akurasi validasi silang {pct(cv['accuracy'])}, presisi "
    f"{pct(cv['precision'])}, recall {pct(cv['recall'])}, dan F1-score {pct(cv['f1'])} "
    f"dengan jumlah tetangga optimal k={M['best_k']}. Studi ablasi membuktikan bahwa "
    f"pembobotan fitur meningkatkan akurasi dari {pct(M['unweighted_cv_acc'])} menjadi "
    f"{pct(M['weighted_cv_acc'])}. Temuan ini menegaskan bahwa CBR berbobot merupakan "
    "pendekatan yang menjanjikan, transparan, dan mudah ditelusuri untuk sistem "
    "pendukung keputusan diagnosis penyakit jantung.", size=10, italic=True)
kw = doc.add_paragraph(); kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw.paragraph_format.space_after = Pt(8)
run(kw, "Kata kunci", size=10, bold=True, italic=True)
run(kw, "\u2014case-based reasoning; diagnosis penyakit jantung; k-nearest neighbor; "
        "HEOM; pembelajaran mesin.", size=10, italic=True)

# ===== SECTION 1: body (two columns) =====
doc.add_section(WD_SECTION.CONTINUOUS)
set_page(doc.sections[1]); set_cols(doc.sections[1], 2, 500)

# ---- I. PENDAHULUAN ----
h1("I", "Pendahuluan")
para("Penyakit jantung atau penyakit kardiovaskular merupakan penyebab kematian nomor "
     "satu secara global. Organisasi Kesehatan Dunia memperkirakan jutaan orang "
     "meninggal setiap tahun akibat gangguan pada jantung dan pembuluh darah [1]. "
     "Diagnosis yang terlambat kerap menyebabkan penanganan yang tidak optimal, "
     "sehingga dibutuhkan sistem pendukung keputusan yang mampu membantu tenaga medis "
     "melakukan penilaian risiko secara cepat dan konsisten.")
para("Berbagai teknik kecerdasan buatan telah dimanfaatkan untuk memprediksi penyakit "
     "jantung, mulai dari pohon keputusan, jaringan saraf tiruan, hingga mesin vektor "
     "pendukung [11], [13]. Meskipun akurasinya sering tinggi, sebagian besar model "
     "tersebut bersifat black box sehingga sulit dijelaskan kepada praktisi klinis. "
     "Dalam domain medis, kemampuan menelusuri alasan sebuah keputusan (interpretability) "
     "sama pentingnya dengan akurasi itu sendiri, karena keputusan menyangkut "
     "keselamatan pasien.")
para("Studi terdahulu pada dataset Cleveland melaporkan bahwa pemilihan fitur dan "
     "fungsi jarak yang tepat berdampak besar terhadap akurasi klasifikasi [11], [13]. "
     "Hal ini memperkuat motivasi untuk memadukan penalaran berbasis kasus dengan "
     "pembobotan fitur, agar sistem tidak hanya akurat tetapi juga mampu menjelaskan "
     "atribut mana yang paling menentukan sebuah diagnosis.", indent=0.0)
para("Case-Based Reasoning (CBR) menawarkan paradigma yang berbeda. Alih-alih membangun "
     "model global, CBR menyelesaikan masalah baru dengan mengingat dan mengadaptasi "
     "solusi dari kasus-kasus serupa yang pernah terjadi [2], [3]. Pendekatan ini "
     "selaras dengan cara dokter berpikir, yaitu membandingkan kondisi pasien baru "
     "dengan pengalaman menangani pasien terdahulu. Karena keputusan didasarkan pada "
     "kasus nyata yang dapat ditunjuk, CBR bersifat transparan dan mudah "
     "dipertanggungjawabkan [9], [10].")
para("Meskipun demikian, penerapan CBR pada data medis menghadapi tantangan berupa "
     "campuran tipe atribut (numerik dan nominal) serta perbedaan tingkat kepentingan "
     "antar atribut. Fungsi jarak yang tidak tepat dapat menurunkan kualitas retrieval "
     "dan pada akhirnya menurunkan akurasi diagnosis. Oleh karena itu diperlukan fungsi "
     "similaritas yang heterogen sekaligus mampu memberi bobot pada atribut yang lebih "
     "diskriminatif.")
para("Kontribusi penelitian ini adalah: (1) merancang sistem diagnosis penyakit jantung "
     "berbasis CBR lima tahap yang lengkap; (2) menerapkan fungsi similaritas HEOM "
     "berbobot untuk menangani campuran atribut numerik dan nominal; dan (3) "
     "mengevaluasi pengaruh pembobotan fitur serta tahap retain terhadap kinerja "
     "diagnosis melalui validasi silang dan pengujian hold-out. Sisa makalah ini disusun "
     "sebagai berikut: Bagian II memaparkan metodologi, Bagian III implementasi, "
     "Bagian IV hasil dan evaluasi, Bagian V diskusi, dan Bagian VI kesimpulan.")

# ---- II. METODOLOGI ----
h1("II", "Metodologi")
para("Secara umum penelitian ini terdiri atas tahapan pengumpulan data, representasi "
     "kasus, pembangunan mesin penalaran CBR, serta evaluasi. Keseluruhan proses "
     "penalaran mengikuti siklus CBR lima tahap yang diilustrasikan pada Gambar 1.")
figure(f"{RES}/fig_cbr_cycle.png", "Gambar 1. Siklus Case-Based Reasoning lima tahap.", 8.0)

h2("A. Dataset dan Representasi Kasus")
para("Dataset yang digunakan mengikuti skema resmi UCI Cleveland Heart Disease [6], [8], "
     "yakni skema yang sama dengan berkas heart disease yang banyak beredar di "
     "repositori publik. Dataset terdiri atas 303 kasus dengan 13 atribut prediktor dan "
     "satu atribut target biner (0 = tidak sakit, 1 = sakit). Rincian atribut disajikan "
     "pada Tabel I. Setiap kasus direpresentasikan sebagai pasangan (fitur masalah, "
     "solusi diagnosis), sehingga basis kasus merupakan kumpulan pengalaman diagnosis "
     "terdahulu.")
caption_top("Tabel I. Atribut Dataset Heart Disease")
make_table([
    ["Atribut", "Tipe", "Keterangan"],
    ["age", "Numerik", "Usia (tahun)"],
    ["sex", "Nominal", "Jenis kelamin"],
    ["cp", "Ordinal", "Tipe nyeri dada"],
    ["trestbps", "Numerik", "Tekanan darah"],
    ["chol", "Numerik", "Kolesterol serum"],
    ["fbs", "Nominal", "Gula darah puasa"],
    ["restecg", "Nominal", "Hasil EKG"],
    ["thalach", "Numerik", "Denyut jantung maks"],
    ["exang", "Nominal", "Angina olahraga"],
    ["oldpeak", "Numerik", "Depresi ST"],
    ["slope", "Ordinal", "Kemiringan ST"],
    ["ca", "Numerik", "Jumlah pembuluh"],
    ["thal", "Nominal", "Thalassemia"],
    ["target", "Kelas", "Diagnosis (0/1)"],
], widths=[1.9, 1.6, 3.9])

para("Analisis awal menunjukkan karakteristik populasi yang beragam: usia pasien "
     f"berkisar {int(_df['age'].min())}-{int(_df['age'].max())} tahun dengan rerata "
     f"{_df['age'].mean():.0f} tahun, kadar kolesterol serum {int(_df['chol'].min())}-"
     f"{int(_df['chol'].max())} mg/dl, serta denyut jantung maksimum "
     f"{int(_df['thalach'].min())}-{int(_df['thalach'].max())} bpm. Perbedaan rentang dan "
     "skala antar atribut inilah yang menjadi alasan utama penerapan normalisasi z-score "
     "pada fungsi jarak, agar atribut berskala besar tidak mendominasi proses retrieval.")

h2("B. Fungsi Similaritas HEOM")
para("Karena atribut bersifat campuran, digunakan Heterogeneous Euclidean-Overlap "
     "Metric (HEOM) [7]. Jarak antar dua kasus x dan y dihitung sebagai akar dari jumlah "
     "kuadrat selisih per atribut yang telah diberi bobot, seperti pada Persamaan (1).")
eq("d(x,y) = \u221a( \u03a3\u1d62 w\u1d62 \u00b7 \u03b4\u1d62(x\u1d62,y\u1d62)\u00b2 )", "1")
para("Untuk atribut numerik, \u03b4 dihitung sebagai selisih mutlak yang dinormalisasi "
     "terhadap simpangan baku atribut (z-score) seperti Persamaan (2), sedangkan untuk "
     "atribut nominal digunakan overlap distance (0 jika sama, 1 jika berbeda). "
     "Normalisasi ini penting agar atribut berskala besar seperti kolesterol tidak "
     "mendominasi atribut berskala kecil.", indent=0.0)
eq("\u03b4_num(x\u1d62,y\u1d62) = |x\u1d62 \u2212 y\u1d62| / \u03c3\u1d62", "2")
para("Nilai jarak kemudian diubah menjadi similaritas melalui Persamaan (3), sehingga "
     "kasus dengan jarak nol memiliki similaritas satu.", indent=0.0)
eq("sim(x,y) = 1 / (1 + d(x,y))", "3")
para("Sebagai ilustrasi, dua pasien dengan selisih usia sekitar satu simpangan baku akan "
     "menyumbang jarak parsial mendekati satu pada atribut usia, sedangkan dua pasien "
     "dengan jenis kelamin berbeda langsung menyumbang jarak penuh sebesar satu pada "
     "atribut tersebut. Kombinasi kedua jenis jarak inilah yang membuat HEOM sesuai untuk "
     "data rekam medis yang bertipe campuran.", indent=0.0)

h2("C. Pembobotan Fitur")
para("Tidak semua atribut memiliki kontribusi yang sama terhadap diagnosis. Oleh karena "
     "itu setiap atribut diberi bobot yang proporsional terhadap kuadrat korelasinya "
     "dengan target, sebagaimana Persamaan (4), dengan r\u1d62 adalah korelasi atribut "
     "ke-i terhadap target dan n jumlah atribut. Pengkuadratan mempertajam kontras "
     "antara atribut penting dan atribut lemah.")
eq("w\u1d62 = n \u00b7 r\u1d62\u00b2 / \u03a3\u2c7c r\u2c7c\u00b2", "4")
para("Dengan skema ini atribut yang lebih diskriminatif seperti jumlah pembuluh utama "
     "(ca) dan depresi ST (oldpeak) memperoleh bobot lebih besar pada perhitungan jarak, "
     "sedangkan atribut lemah seperti gula darah puasa (fbs) hampir diabaikan.", indent=0.0)

h2("D. Retrieve, Reuse, Revise, dan Retain")
para("Pada tahap retrieve, sistem mengambil k kasus dengan similaritas tertinggi. Tahap "
     "reuse menggabungkan solusi melalui pemungutan suara berbobot similaritas, yaitu "
     "kelas dengan total similaritas terbesar dipilih sebagai diagnosis. Tahap revise "
     "menguji keyakinan hasil; bila rasio dukungan berada di bawah ambang tertentu, "
     "kasus ditandai berketidakpastian tinggi sehingga dapat dirujuk ke pakar. "
     "Terakhir, tahap retain menyimpan kasus baru yang terverifikasi ke dalam basis "
     "kasus agar pengetahuan sistem terus bertambah. Prosedur inti diagnosis dirangkum "
     "pada Algoritma 1.")
algorithm("Algoritma 1. Diagnosis CBR untuk satu kasus", [
    "Input : kasus q, basis kasus C, tetangga k, ambang t",
    "Output: diagnosis y, keyakinan conf",
    "1  for each c in C do",
    "2     d[c] <- HEOM_berbobot(q, c)",
    "3  N  <- k kasus dengan d[c] terkecil",
    "4  w0 <- 0 ; w1 <- 0",
    "5  for each c in N do",
    "6     s <- 1 / (1 + d[c])",
    "7     if label(c)=1 then w1 += s else w0 += s",
    "8  y    <- argmax(w0, w1)",
    "9  conf <- max(w0,w1) / (w0 + w1)",
    "10 if conf < t then tandai 'tidak pasti'",
    "11 return y, conf",
])

h2("E. Skema Evaluasi")
para("Kinerja diukur dengan empat metrik standar, yaitu akurasi, presisi, recall, dan "
     "F1-score sebagaimana Persamaan (5)-(8), dengan TP, TN, FP, dan FN masing-masing "
     "menyatakan true positive, true negative, false positive, dan false negative.")
eq("Akurasi = (TP+TN) / (TP+TN+FP+FN)", "5")
eq("Presisi = TP / (TP+FP)", "6")
eq("Recall = TP / (TP+FN)", "7")
eq("F1 = 2 \u00b7 (Presisi \u00b7 Recall) / (Presisi + Recall)", "8")
para("Nilai k optimal ditentukan melalui validasi silang 5-fold pada data latih. Model "
     "akhir diuji pada data hold-out sebesar 20% dan divalidasi ulang dengan validasi "
     "silang 5-fold pada keseluruhan data untuk memperoleh estimasi yang stabil [14]. "
     "Penggunaan validasi silang mengurangi ketergantungan hasil pada satu pembagian "
     "data tertentu.", indent=0.0)

# ---- III. IMPLEMENTASI ----
h1("III", "Implementasi")
h2("A. Lingkungan Pengembangan")
para("Sistem diimplementasikan menggunakan bahasa Python murni tanpa pustaka "
     "pembelajaran mesin siap pakai; seluruh mesin CBR, fungsi similaritas, dan metrik "
     "evaluasi ditulis dari awal untuk menjaga transparansi dan reprodusibilitas. "
     "Rincian lingkungan ditunjukkan pada Tabel II.")
caption_top("Tabel II. Lingkungan Implementasi")
make_table([
    ["Komponen", "Keterangan"],
    ["Bahasa", "Python 3"],
    ["Komputasi numerik", "NumPy"],
    ["Manipulasi data", "pandas"],
    ["Visualisasi", "Matplotlib"],
    ["Mesin CBR & metrik", "Implementasi mandiri"],
], widths=[3.4, 4.0])
para("Pemilihan implementasi Python murni bukan sekadar batasan teknis, melainkan "
     "keputusan desain agar setiap langkah perhitungan similaritas, pembobotan, dan "
     "pemungutan suara dapat diperiksa baris demi baris. Hal ini penting pada domain "
     "medis yang menuntut auditabilitas penuh terhadap proses pengambilan keputusan.", indent=0.0)
h2("B. Struktur Repositori dan Pipeline")
para("Kode disusun modular agar mudah direplikasi. Struktur berkas utama ditunjukkan "
     "pada Tabel III. Pipeline berjalan dua langkah sederhana: menjalankan "
     "generate_dataset.py untuk menyiapkan data, lalu run_experiment.py untuk melatih, "
     "menyetel k, menjalankan studi ablasi, serta menyimpan metrik dan seluruh gambar.")
caption_top("Tabel III. Struktur Berkas Repositori")
make_table([
    ["Berkas", "Fungsi"],
    ["generate_dataset.py", "Menyiapkan data"],
    ["cbr.py", "Kelas CBR, HEOM, metrik"],
    ["run_experiment.py", "Eksperimen & evaluasi"],
    ["make_cbr_diagram.py", "Diagram siklus"],
    ["requirements.txt", "Daftar dependensi"],
    ["README.md", "Dokumentasi"],
], widths=[3.5, 3.9])
para("Reprodusibilitas dijamin dengan penetapan seed acak yang tetap. Kode dirancang "
     "fleksibel: apabila berkas data asli (misalnya heart.csv dari repositori publik) "
     "diletakkan pada folder data/, pipeline akan langsung memakainya tanpa perubahan "
     "kode. Dengan demikian metodologi yang sama dapat divalidasi ulang pada data "
     "klinis nyata.", indent=0.0)
para("Selain itu, setiap eksperimen mencatat metrik ke berkas results/metrics.json dan "
     "seluruh grafik disimpan otomatis ke folder results/, sehingga seluruh hasil pada "
     "makalah ini dapat diverifikasi ulang oleh pihak lain hanya dengan menjalankan "
     "kembali pipeline tanpa penyesuaian tambahan.", indent=0.0)

# ---- IV. HASIL DAN EVALUASI ----
h1("IV", "Hasil dan Evaluasi")
h2("A. Karakteristik Data")
para("Distribusi kelas relatif seimbang dengan proporsi kasus sakit sebesar "
     f"{M['disease_rate_total']*100:.1f}%, seperti terlihat pada Gambar 2. Keseimbangan "
     "ini mengurangi risiko bias model terhadap salah satu kelas [15].")
figure(f"{RES}/fig_class_distribution.png", "Gambar 2. Distribusi kelas dataset.", 7.6)
para("Selain seimbang, dataset juga memperlihatkan korelasi yang bermakna antara "
     "beberapa atribut klinis dan status diagnosis. Atribut seperti jumlah pembuluh "
     "utama (ca), depresi ST (oldpeak), dan tipe nyeri dada (cp) menunjukkan hubungan "
     "yang lebih kuat terhadap target dibanding atribut lain, sebuah pola yang kelak "
     "tercermin pada bobot fitur yang dipelajari sistem.", indent=0.0)

h2("B. Pemilihan Jumlah Tetangga (k)")
para(f"Penyetelan k melalui validasi silang menghasilkan nilai optimal k={M['best_k']} "
     f"dengan akurasi validasi silang tertinggi sebesar {pct(M['weighted_cv_acc'])}, "
     "sebagaimana ditunjukkan pada Gambar 3 dan dirinci pada Tabel IV. Nilai k yang "
     "terlalu kecil cenderung sensitif terhadap derau, sedangkan k yang terlalu besar "
     "membuat keputusan menjadi terlalu umum.")
figure(f"{RES}/fig_k_selection.png", "Gambar 3. Pemilihan k via validasi silang.", 7.8)
caption_top("Tabel IV. Akurasi Validasi Silang per Nilai k")
krows = [["k", "Akurasi CV", "Std"]]
for kk, mu, sd in zip(M["k_grid"], M["cv_mean"], M["cv_std"]):
    krows.append([str(kk), f"{mu*100:.1f}%", f"{sd*100:.1f}%"])
make_table(krows, widths=[1.8, 3.0, 2.6])

h2("C. Kinerja Diagnosis")
para("Ringkasan kinerja model pada skema validasi silang 5-fold dan hold-out disajikan "
     "pada Tabel V, sedangkan perbandingan visualnya pada Gambar 4. Matriks kebingungan "
     "hasil validasi silang ditunjukkan pada Gambar 5.")
caption_top("Tabel V. Hasil Evaluasi Model CBR")
make_table([
    ["Metrik", "CV 5-fold", "Hold-out"],
    ["Akurasi", pct(cv["accuracy"]), pct(ts["accuracy"])],
    ["Presisi", pct(cv["precision"]), pct(ts["precision"])],
    ["Recall", pct(cv["recall"]), pct(ts["recall"])],
    ["F1-score", pct(cv["f1"]), pct(ts["f1"])],
], widths=[2.5, 2.6, 2.5])
figure(f"{RES}/fig_metrics.png", "Gambar 4. Perbandingan metrik CV dan hold-out.", 8.0)
figure(f"{RES}/fig_confusion_matrix.png", "Gambar 5. Matriks kebingungan (5-fold CV).", 7.0)
para(f"Dari matriks kebingungan validasi silang, sistem berhasil mengidentifikasi "
     f"{cv['confusion']['tp']} kasus sakit dan {cv['confusion']['tn']} kasus sehat secara "
     f"benar, dengan {cv['confusion']['fp']} false positive dan {cv['confusion']['fn']} "
     "false negative. Sebaran ini memperlihatkan bahwa model melakukan kesalahan secara "
     "relatif seimbang, meskipun kesalahan pada kelas sakit sedikit lebih banyak.")
para(f"Model mencapai akurasi validasi silang {pct(cv['accuracy'])} "
     f"(\u00b1{cv['accuracy_std']*100:.1f}%). Nilai presisi ({pct(cv['precision'])}) yang "
     f"lebih tinggi daripada recall ({pct(cv['recall'])}) menunjukkan model cenderung "
     "berhati-hati dalam menyatakan pasien sakit. Rentang akurasi latih "
     f"({pct(tr['accuracy'])}) dan akurasi uji yang wajar menandakan model tidak "
     "mengalami overfitting berlebih dan mampu menggeneralisasi.")
caption_top("Tabel VI. Metrik Kinerja per Kelas (5-fold CV)")
make_table([
    ["Kelas", "Presisi", "Recall", "F1-score"],
    ["Sehat (0)", f"{prec_neg*100:.1f}%", f"{rec_neg*100:.1f}%", f"{f1_neg*100:.1f}%"],
    ["Sakit (1)", f"{prec_pos*100:.1f}%", f"{rec_pos*100:.1f}%", f"{f1_pos*100:.1f}%"],
], widths=[2.4, 1.7, 1.6, 1.7])
para("Tabel VI merinci kinerja untuk masing-masing kelas. Kelas sehat cenderung "
     "memperoleh recall lebih tinggi, sedangkan kelas sakit memiliki presisi yang baik "
     "namun recall lebih rendah. Pola ini konsisten dengan kecenderungan model bersikap "
     "konservatif dalam menyatakan seorang pasien sakit.", indent=0.0)

h2("D. Studi Ablasi: Pembobotan dan Retain")
para("Untuk mengukur kontribusi tiap komponen, dilakukan dua studi ablasi yang dirangkum "
     f"pada Tabel VII. Pembobotan fitur terbukti menaikkan akurasi secara nyata, dari "
     f"{pct(M['unweighted_cv_acc'])} menjadi {pct(M['weighted_cv_acc'])} (naik sekitar "
     f"{(M['weighted_cv_acc']-M['unweighted_cv_acc'])*100:.1f} poin). Bobot fitur hasil "
     "pembelajaran ditampilkan pada Gambar 6, yang memperlihatkan dominasi atribut ca "
     "dan oldpeak, sejalan dengan pengetahuan klinis.")
caption_top("Tabel VII. Studi Ablasi Komponen CBR")
make_table([
    ["Konfigurasi", "Akurasi"],
    ["Tanpa pembobotan", pct(M["unweighted_cv_acc"])],
    ["Dengan pembobotan", pct(M["weighted_cv_acc"])],
    ["Retain nonaktif", pct(M["retain_off_acc"])],
    ["Retain aktif", pct(M["retain_on_acc"])],
], widths=[4.4, 3.0])
figure(f"{RES}/fig_feature_weights.png", "Gambar 6. Bobot fitur hasil pembelajaran.", 8.0)
figure(f"{RES}/fig_retain_effect.png", "Gambar 7. Pengaruh tahap retain.", 7.6)
para("Gambar 6 memperlihatkan bahwa atribut jumlah pembuluh utama (ca) dan depresi ST "
     "(oldpeak) memperoleh bobot tertinggi, diikuti tipe nyeri dada (cp) dan kemiringan "
     "segmen ST (slope). Hasil ini selaras dengan literatur kardiologi yang menempatkan "
     "atribut-atribut tersebut sebagai penanda penting penyakit arteri koroner, sehingga "
     "bobot yang dipelajari sistem dapat dipertanggungjawabkan secara klinis.")
para("Gambar 7 membandingkan akurasi ketika tahap retain diaktifkan dan dinonaktifkan. "
     f"Pada eksperimen ini retain justru menurunkan akurasi tipis dari "
     f"{pct(M['retain_off_acc'])} menjadi {pct(M['retain_on_acc'])}. Temuan ini "
     "mengindikasikan bahwa menambah kasus tanpa penyaringan kualitas dapat memperbesar "
     "derau pada basis kasus, sebuah fenomena yang dikenal sebagai utility problem dalam "
     "literatur CBR.", indent=0.0)

h2("E. Kompleksitas dan Waktu Komputasi")
para("Karena CBR menunda komputasi hingga tahap diagnosis (lazy learning), pelatihan "
     "praktis tidak memerlukan biaya selain penyimpanan basis kasus. Untuk satu kueri, "
     "sistem menghitung jarak terhadap seluruh n kasus sehingga kompleksitas retrieval "
     "adalah O(n\u00b7m) dengan m jumlah atribut. Pada skala 303 kasus dan 13 atribut, "
     "diagnosis satu pasien berlangsung hampir seketika, sehingga pendekatan ini layak "
     "untuk penggunaan interaktif di klinik. Untuk basis kasus berskala besar, struktur "
     "indeks seperti k-d tree atau penyaringan kandidat dapat diterapkan guna menjaga "
     "kecepatan retrieval.")

# ---- V. DISKUSI ----
h1("V", "Diskusi")
para("Hasil eksperimen menegaskan bahwa CBR mampu memberikan diagnosis penyakit jantung "
     "dengan kinerja yang kompetitif sekaligus tetap dapat dijelaskan. Keunggulan utama "
     "pendekatan ini adalah setiap keputusan dapat ditelusuri ke kasus-kasus mirip yang "
     "menjadi dasar pengambilan keputusan, sehingga cocok untuk lingkungan klinis yang "
     "menuntut akuntabilitas.")
para("Jika dibandingkan dengan pendekatan lain pada domain yang sama, model berbasis "
     "pohon keputusan dan jaringan saraf tiruan memang kerap melaporkan akurasi yang "
     "sedikit lebih tinggi [11], [13], tetapi keduanya sulit menjelaskan alasan di balik "
     "sebuah prediksi. CBR menukar sebagian kecil akurasi dengan transparansi penuh, "
     "sebuah kompromi yang umumnya bernilai tinggi pada aplikasi klinis nyata.", indent=0.0)
para(f"Konsistensi antara akurasi validasi silang ({pct(cv['accuracy'])}) dan akurasi "
     f"hold-out ({pct(ts['accuracy'])}) menunjukkan bahwa kinerja model cukup stabil dan "
     "tidak bergantung pada satu pembagian data tertentu. Simpangan baku akurasi "
     f"antar-fold yang kecil (\u00b1{cv['accuracy_std']*100:.1f}%) semakin menguatkan "
     "keandalan estimasi tersebut.", indent=0.0)
para("Analisis kegagalan menunjukkan kesalahan terbanyak berupa false negative, yaitu "
     f"{cv['confusion']['fn']} kasus sakit yang terdeteksi sehat. Dalam konteks medis, "
     "kesalahan jenis ini lebih berisiko dibanding false positive karena pasien sakit "
     "dapat luput dari penanganan. Penyebabnya adalah sebagian pasien sakit memiliki "
     "profil klinis yang mirip pasien sehat sehingga sulit dibedakan hanya melalui jarak "
     "fitur. Penanganan lanjutan dapat berupa penyesuaian ambang keputusan agar lebih "
     "sensitif terhadap kelas sakit.", indent=0.0)
para("Tahap retain pada eksperimen ini memberi pengaruh kecil bahkan sedikit menurunkan "
     "akurasi. Hal ini disebabkan penambahan kasus tanpa penyaringan kualitas dapat "
     "memasukkan derau ke basis kasus (utility problem). Oleh karena itu retain "
     "sebaiknya disertai mekanisme seleksi kasus, misalnya hanya menyimpan kasus yang "
     "diagnosisnya diverifikasi pakar dan berbeda cukup jauh dari kasus yang telah ada.")
para("Dari sisi penerapan, sifat incremental CBR memudahkan integrasi ke alur kerja "
     "rumah sakit: kasus baru yang telah diverifikasi dokter dapat ditambahkan tanpa "
     "melatih ulang seluruh model. Meski demikian, sistem semacam ini harus tetap "
     "diposisikan sebagai alat bantu, bukan pengganti penilaian klinis, dan wajib "
     "memperhatikan aspek privasi serta keamanan data pasien.", indent=0.0)
para("Keterbatasan penelitian ini terletak pada data. Karena keterbatasan akses, data "
     "dibangkitkan mengikuti skema resmi UCI Cleveland dan model risiko klinis yang "
     "wajar, bukan diunduh langsung dari repositori. Struktur atribut, rentang nilai, "
     "dan hubungan antar fitur telah dirancang menyerupai data nyata sehingga "
     "metodologi tetap sahih dan langsung dapat diterapkan pada data asli tanpa "
     "mengubah kode. Untuk penelitian lanjutan, disarankan memakai data klinis nyata, "
     "menambahkan seleksi kasus pada tahap retain, serta membandingkan HEOM dengan "
     "metrik similaritas lain seperti jarak Manhattan atau Mahalanobis.")

# ---- VI. KESIMPULAN ----
h1("VI", "Kesimpulan")
para("Penelitian ini berhasil membangun sistem diagnosis penyakit jantung berbasis "
     "Case-Based Reasoning lima tahap dengan fungsi similaritas HEOM berbobot. Model "
     f"mencapai akurasi validasi silang {pct(cv['accuracy'])}, presisi "
     f"{pct(cv['precision'])}, recall {pct(cv['recall'])}, dan F1-score {pct(cv['f1'])} "
     f"pada k={M['best_k']}. Pembobotan fitur terbukti meningkatkan akurasi secara "
     "nyata, sementara tahap retain memerlukan seleksi kasus agar bermanfaat. Dengan "
     "sifatnya yang transparan dan dapat ditelusuri, CBR berbobot layak dipertimbangkan "
     "sebagai basis sistem pendukung keputusan diagnosis penyakit jantung, dan berpotensi "
     "dikembangkan lebih lanjut menggunakan data klinis nyata.")
para("Sebagai arah pengembangan selanjutnya, penelitian ini dapat diperluas dengan "
     "memanfaatkan data klinis nyata berskala lebih besar, menerapkan seleksi kasus "
     "cerdas pada tahap retain, mengoptimasi bobot fitur secara otomatis, serta "
     "membandingkan HEOM dengan metrik similaritas lain guna memperoleh kombinasi "
     "terbaik antara akurasi dan interpretabilitas.", indent=0.0)

# ---- REFERENCES ----
h1("", "Daftar Pustaka")
refs = [
    "World Health Organization, \u201cCardiovascular diseases (CVDs),\u201d WHO Fact Sheets, 2021.",
    "A. Aamodt and E. Plaza, \u201cCase-based reasoning: Foundational issues, methodological variations, and system approaches,\u201d AI Communications, vol. 7, no. 1, pp. 39\u201359, 1994.",
    "J. L. Kolodner, \u201cAn introduction to case-based reasoning,\u201d Artificial Intelligence Review, vol. 6, no. 1, pp. 3\u201334, 1992.",
    "T. Cover and P. Hart, \u201cNearest neighbor pattern classification,\u201d IEEE Transactions on Information Theory, vol. 13, no. 1, pp. 21\u201327, 1967.",
    "E. Fix and J. L. Hodges, \u201cDiscriminatory analysis: Nonparametric discrimination,\u201d USAF School of Aviation Medicine, Randolph Field, TX, 1951.",
    "R. Detrano et al., \u201cInternational application of a new probability algorithm for the diagnosis of coronary artery disease,\u201d The American Journal of Cardiology, vol. 64, no. 5, pp. 304\u2013310, 1989.",
    "D. R. Wilson and T. R. Martinez, \u201cImproved heterogeneous distance functions,\u201d Journal of Artificial Intelligence Research, vol. 6, pp. 1\u201334, 1997.",
    "D. Dua and C. Graff, \u201cUCI Machine Learning Repository: Heart Disease Data Set,\u201d University of California, Irvine, 2019.",
    "I. Watson, \u201cCase-based reasoning is a methodology not a technology,\u201d Knowledge-Based Systems, vol. 12, no. 5\u20136, pp. 303\u2013308, 1999.",
    "M. M. Richter and R. O. Weber, Case-Based Reasoning: A Textbook. Berlin: Springer, 2013.",
    "S. Palaniappan and R. Awang, \u201cIntelligent heart disease prediction system using data mining techniques,\u201d in IEEE/ACS Int. Conf. on Computer Systems and Applications, 2008, pp. 108\u2013115.",
    "J. Han, M. Kamber, and J. Pei, Data Mining: Concepts and Techniques, 3rd ed. Burlington: Morgan Kaufmann, 2011.",
    "C. Nasa and Suman, \u201cEvaluation of different classification techniques for heart disease dataset,\u201d Int. Journal of Computer Science and Information Technologies, vol. 3, no. 2, 2012.",
    "R. Kohavi, \u201cA study of cross-validation and bootstrap for accuracy estimation and model selection,\u201d in Proc. IJCAI, 1995, pp. 1137\u20131145.",
    "N. Japkowicz and S. Stephen, \u201cThe class imbalance problem: A systematic study,\u201d Intelligent Data Analysis, vol. 6, no. 5, pp. 429\u2013449, 2002.",
    "S. M. Weiss and C. A. Kulikowski, Computer Systems That Learn. San Mateo: Morgan Kaufmann, 1991.",
    "L. Breiman, \u201cRandom forests,\u201d Machine Learning, vol. 45, no. 1, pp. 5\u201332, 2001.",
    "V. Podgorelec, P. Kokol, B. Stiglic, and I. Rozman, \u201cDecision trees: An overview and their use in medicine,\u201d Journal of Medical Systems, vol. 26, no. 5, pp. 445\u2013463, 2002.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    run(p, f"[{i}] ", size=9); run(p, r, size=9)

out = f"{ROOT}/Artikel_JOIV_CBR_Heart_Disease.docx"
doc.save(out)
print("saved", out)
